# -*- coding: utf-8 -*-
"""
Document Service
Handles document ingestion and processing
"""
import logging
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from collections import Counter
from config import get_config

logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("pypdf not available - PDF extraction disabled")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX extraction disabled")

# File categories for storage analysis
FILE_CATEGORIES = {
    # Audio
    "mp3": "Muziek", "wav": "Muziek", "flac": "Muziek",
    # Video
    "mp4": "Video's", "avi": "Video's", "mkv": "Video's",
    # Images
    "jpg": "Foto's", "jpeg": "Foto's", "png": "Foto's",
    "gif": "Foto's", "webp": "Foto's",
    # E-books
    "epub": "E-books", "pdf": "E-books", "mobi": "E-books",
    # Documents
    "docx": "Documenten", "doc": "Documenten",
    "xlsx": "Documenten", "xls": "Documenten",
    "pptx": "Documenten", "txt": "Documenten",
}

# Document types suitable for RAG ingestion
RAG_DOCUMENT_TYPES = {
    "txt", "pdf", "docx", "doc", "md", "rst",
    "epub", "mobi", "html", "xml", "json", "csv"
}

class DocumentService:
    """Service for document operations"""
    
    def __init__(self, config=None):
        """Initialize document service"""
        if config is None:
            config = get_config()
        
        self.config = config
        self.chunk_size = config.MAX_CHUNK_SIZE
        self.chunk_overlap = config.CHUNK_OVERLAP
        logger.info("Document service initialized")
    
    def scan_directory(self, directory: Path, excluded_dirs: set = None) -> List[Dict]:
        """Scan directory for supported documents
        
        Args:
            directory: Path to directory to scan
            excluded_dirs: Set of directory names to exclude
            
        Returns:
            List of document dictionaries with metadata
        """
        if excluded_dirs is None:
            excluded_dirs = {'AppData', 'node_modules', '.git', '__pycache__'}
        
        documents = []
        file_count = 0
        
        logger.info(f"Scanning directory: {directory}")
        
        try:
            for file_path in directory.rglob("*"):
                try:
                    # Skip excluded directories
                    if any(excluded in file_path.parts for excluded in excluded_dirs):
                        continue
                    
                    if not file_path.is_file():
                        continue
                    
                    # Get file stats
                    try:
                        file_stat = file_path.stat()
                        file_size = file_stat.st_size
                    except (OSError, FileNotFoundError):
                        continue
                    
                    # Check if suitable for RAG
                    ext = file_path.suffix.lower().lstrip('.')
                    if ext in RAG_DOCUMENT_TYPES and file_size < self.config.MAX_DOCUMENT_SIZE:
                        documents.append({
                            'path': str(file_path),
                            'name': file_path.name,
                            'type': ext,
                            'size': file_size,
                            'size_formatted': self._format_size(file_size)
                        })
                    
                    file_count += 1
                    
                    if file_count % 1000 == 0:
                        logger.info(f"Processed {file_count} files...")
                        
                except (PermissionError, OSError):
                    continue
                    
        except KeyboardInterrupt:
            logger.warning("Scan interrupted by user")
        
        logger.info(f"Found {len(documents)} documents in {file_count} files")
        return documents
    
    def ingest_documents(self, file_paths: List[Path]) -> int:
        """Ingest multiple documents
        
        Args:
            file_paths: List of document paths to ingest
            
        Returns:
            Number of successfully ingested documents
        """
        from services import DatabaseService, LLMService
        
        db_service = DatabaseService(self.config)
        llm_service = LLMService(self.config)
        
        ingested_count = 0
        
        for i, file_path in enumerate(file_paths, 1):
            logger.info(f"Ingesting [{i}/{len(file_paths)}]: {file_path.name}")
            
            try:
                # Extract text
                text = self.extract_text(file_path)
                if not text or len(text.strip()) < 50:
                    logger.warning(f"Insufficient content: {file_path.name}")
                    continue
                
                # Chunk text
                chunks = self.chunk_text(text)
                logger.info(f"Created {len(chunks)} chunks")
                
                # Generate embeddings
                embeddings = []
                for chunk in chunks:
                    embedding = llm_service.generate_embedding(chunk)
                    if not embedding:
                        logger.error(f"Failed to generate embedding for {file_path.name}")
                        break
                    embeddings.append(embedding)
                
                if len(embeddings) != len(chunks):
                    logger.error(f"Embedding generation incomplete for {file_path.name}")
                    continue
                
                # Store in database
                file_stat = file_path.stat()
                success = db_service.ingest_document(
                    file_path=file_path,
                    chunks=chunks,
                    embeddings=embeddings,
                    modified_time=file_stat.st_mtime
                )
                
                if success:
                    ingested_count += 1
                    logger.info(f"Successfully ingested: {file_path.name}")
                    
            except Exception as e:
                logger.error(f"Error ingesting {file_path.name}: {e}")
                continue
        
        logger.info(f"Ingested {ingested_count}/{len(file_paths)} documents")
        return ingested_count
    
    def extract_text(self, file_path: Path) -> Optional[str]:
        """Extract text from various file types with improved DOCX handling
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text or None
        """
        ext = file_path.suffix.lower().lstrip('.')
        
        try:
            # Plain text files
            if ext in ["txt", "md", "rst", "csv", "json", "xml", "html"]:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                    logger.debug(f"Extracted {len(text)} chars from {file_path.name}")
                    return text
            
            # PDF files
            elif ext == "pdf" and PYPDF_AVAILABLE:
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    text_parts = []
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    text = "\n".join(text_parts)
                    logger.debug(f"Extracted {len(text)} chars from {len(reader.pages)} pages in {file_path.name}")
                    return text
            
            # Word documents - Enhanced extraction
            elif ext == "docx" and DOCX_AVAILABLE:
                doc = docx.Document(file_path)
                text_parts = []
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)
                
                # Extract headers/footers
                for section in doc.sections:
                    # Header
                    header = section.header
                    for para in header.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    
                    # Footer
                    footer = section.footer
                    for para in footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                
                text = "\n".join(text_parts)
                logger.info(f"Extracted {len(text)} chars from DOCX {file_path.name} (paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)})")
                return text
            
            return None
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        separators = ["\n\n", "\n", ". ", " ", ""]
        
        def split_recursive(text: str, sep_index: int = 0) -> List[str]:
            if len(text) <= self.chunk_size:
                return [text] if text.strip() else []
            
            if sep_index >= len(separators):
                # Character-level split with overlap
                result = []
                start = 0
                while start < len(text):
                    end = min(start + self.chunk_size, len(text))
                    chunk = text[start:end]
                    if chunk.strip():
                        result.append(chunk)
                    if end >= len(text):
                        break
                    start = end - self.chunk_overlap
                return result
            
            separator = separators[sep_index]
            splits = text.split(separator)
            
            result = []
            current_chunk = ""
            
            for i, split in enumerate(splits):
                piece = split + (separator if i < len(splits) - 1 else "")
                
                if len(current_chunk) + len(piece) <= self.chunk_size:
                    current_chunk += piece
                else:
                    if current_chunk.strip():
                        result.append(current_chunk)
                    
                    if len(piece) > self.chunk_size:
                        result.extend(split_recursive(piece, sep_index + 1))
                        current_chunk = ""
                    else:
                        if result:
                            prev_chunk = result[-1]
                            overlap_text = prev_chunk[-self.chunk_overlap:] if len(prev_chunk) > self.chunk_overlap else prev_chunk
                            current_chunk = overlap_text + piece
                        else:
                            current_chunk = piece
            
            if current_chunk.strip():
                result.append(current_chunk)
            
            return result
        
        chunks = split_recursive(text)
        chunks = [c.strip() for c in chunks if c.strip()]
        
        return chunks if chunks else [text]
    
    def analyze_storage(self, directory: Path) -> Dict:
        """Analyze storage distribution in directory
        
        Args:
            directory: Path to analyze
            
        Returns:
            Dictionary with storage analysis data
        """
        categories = Counter()
        total_size = 0
        file_count = 0
        
        logger.info(f"Analyzing storage: {directory}")
        
        try:
            for file_path in directory.rglob("*"):
                try:
                    if not file_path.is_file():
                        continue
                    
                    file_stat = file_path.stat()
                    file_size = file_stat.st_size
                    
                    ext = file_path.suffix.lower().lstrip('.')
                    category = FILE_CATEGORIES.get(ext, "Overig")
                    
                    categories[category] += file_size
                    total_size += file_size
                    file_count += 1
                    
                except (PermissionError, OSError):
                    continue
                    
        except KeyboardInterrupt:
            logger.warning("Analysis interrupted")
        
        # Format results
        results = {
            'total_size': total_size,
            'total_size_formatted': self._format_size(total_size),
            'file_count': file_count,
            'categories': [
                {
                    'name': cat,
                    'size': size,
                    'size_formatted': self._format_size(size),
                    'percentage': round((size / total_size * 100), 1) if total_size > 0 else 0
                }
                for cat, size in categories.most_common()
            ]
        }
        
        logger.info(f"Analysis complete: {file_count} files, {self._format_size(total_size)}")
        return results
    
    def _format_size(self, bytes_size: int) -> str:
        """Format bytes to human-readable size"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if bytes_size < 1024:
                return f"{bytes_size:.2f} {unit}"
            bytes_size /= 1024
        return f"{bytes_size:.2f} PB"
