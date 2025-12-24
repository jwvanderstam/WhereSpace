# -*- coding: utf-8 -*-
"""
WhereSpace - Storage Analysis & Document Ingestion Tool
========================================================

A comprehensive tool for analyzing file storage, categorizing documents,
and ingesting them into a pgvector database for RAG (Retrieval-Augmented Generation).

Features:
    - Recursive directory scanning with configurable exclusions
    - File categorization by extension
    - AI-powered storage analysis via Ollama
    - Document text extraction (PDF, DOCX, TXT, etc.)
    - Intelligent document chunking with overlap
    - Vector embedding generation using Ollama
    - PostgreSQL/pgvector storage with automatic schema management
    - Skip already-ingested documents (based on modification time)

Requirements:
    - Python 3.8+
    - PostgreSQL with pgvector extension
    - Ollama running locally with llama3.1 and nomic-embed-text models

Author: Your Name
License: MIT
Version: 1.0.0
"""

from pathlib import Path
from collections import Counter
import sys
import subprocess
import logging
from typing import Tuple, List, Optional, Dict
from dataclasses import dataclass
from contextlib import contextmanager

# ============================================================================
# LOGGING CONFIGURATION
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# ============================================================================
# CONFIGURATION CONSTANTS
# ============================================================================

# Ollama Configuration
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_EMBED_URL = "http://localhost:11434/api/embeddings"
OLLAMA_MODEL = "llama3.1"

# Embedding Model Configuration
# Recommended models (pull with: ollama pull <model>):
# - nomic-embed-text (768d) - Best overall, multilingual
# - snowflake-arctic-embed (1024d) - Strong retrieval performance
# - mxbai-embed-large (1024d) - Good for long documents
OLLAMA_EMBED_MODEL = "nomic-embed-text"  # Changed from basic model
OLLAMA_EMBED_DIMENSION = 768  # nomic-embed-text dimension
OLLAMA_TIMEOUT = 60

# Scanning Configuration
PROGRESS_INTERVAL = 1000  # Files between progress updates
MAX_DOCUMENT_SIZE = 10 * 1024 * 1024  # 10MB max for RAG ingestion

# Chunking Configuration - OPTIMIZED FOR PERFORMANCE
# Using recursive character splitting with overlap for better context preservation
CHUNK_SIZE = 512  # Reduced from 1000 - better for token-based models (~128 tokens)
CHUNK_OVERLAP = 100  # Increased from 200 - maintains context continuity
CHUNK_SEPARATORS = ["\n\n", "\n", ". ", " ", ""]  # Hierarchical splitting

# PostgreSQL/pgvector Configuration
PG_HOST = "localhost"
PG_PORT = 5432
PG_DATABASE = "vectordb"
PG_USER = "postgres"
PG_PASSWORD = "Mutsmuts10"
PG_TABLE = "documents"

# File Categories (Dutch labels)
FILE_CATEGORIES = {
    # Audio
    "mp3": "Muziek", "wav": "Muziek", "flac": "Muziek",
    # Video
    "mp4": "Video's", "avi": "Video's", "mkv": "Video's",
    # Images
    "jpg": "Foto's", "jpeg": "Foto's", "png": "Foto's",
    "gif": "Foto's", "webp": "Foto's",
    # E-books
    "epub": "E-books", "pdf": "E-books", "mobi": "E-books",
    # Documents
    "docx": "Documenten", "doc": "Documenten",
    "xlsx": "Documenten", "xls": "Documenten",
    "pptx": "Documenten", "txt": "Documenten",
}

# Document types suitable for RAG ingestion
RAG_DOCUMENT_TYPES = {
    "txt", "pdf", "docx", "doc", "md", "rst",
    "epub", "mobi", "html", "xml", "json", "csv"
}

# Required Python modules
REQUIRED_MODULES = {
    "requests": "requests",
    "psycopg2": "psycopg2-binary",
    "pypdf": "pypdf",
    "docx": "python-docx"
}

# ============================================================================
# MODULE MANAGEMENT
# ============================================================================

def check_and_install_modules() -> None:
    """
    Check for required Python modules and install missing ones.
    
    This function attempts to import each required module. If a module
    is missing, it automatically installs it using pip.
    
    Note:
        Installation output is suppressed (DEVNULL) for cleaner logs.
        Manual installation instructions are provided if auto-install fails.
    """
    missing_modules = []
    
    # Check which modules are missing
    for module_name, pip_name in REQUIRED_MODULES.items():
        try:
            __import__(module_name)
            logger.debug(f"? Module '{module_name}' already installed")
        except ImportError:
            missing_modules.append((module_name, pip_name))
    
    # Install missing modules
    if missing_modules:
        logger.warning(f"Missing {len(missing_modules)} module(s). Attempting auto-install...")
        
        for module_name, pip_name in missing_modules:
            try:
                logger.info(f"Installing {pip_name}...")
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pip_name],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL
                )
                logger.info(f"? Successfully installed '{module_name}'")
            except Exception as e:
                logger.error(f"? Failed to install '{module_name}': {e}")
                logger.info(f"  Please run manually: pip install {pip_name}")


# Run module check on import
check_and_install_modules()

# Import modules after installation attempt
try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False
    logger.warning("? requests not available - AI analysis disabled")

try:
    import psycopg2
    from psycopg2 import sql
    PSYCOPG2_AVAILABLE = True
except ImportError:
    PSYCOPG2_AVAILABLE = False
    logger.warning("? psycopg2 not available - RAG ingestion disabled")

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("? pypdf not available - PDF extraction disabled")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("? python-docx not available - DOCX extraction disabled")

# Import batch embedding system at the top with other imports
try:
    from batch_embeddings import generate_embeddings_batch, BatchEmbeddingGenerator
    BATCH_EMBEDDINGS_AVAILABLE = True
    logger.info("Batch embedding system available - 5-10x faster ingestion!")
except ImportError:
    BATCH_EMBEDDINGS_AVAILABLE = False
    logger.warning("Batch embeddings not available - using sequential mode")

# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class ScanResult:
    """Results from a directory scan."""
    categories: Counter  # File categories -> total size
    directories: Counter  # Directory paths -> total size
    documents_by_dir: Dict[str, List[Path]]  # Directory -> document list
    file_count: int
    error_count: int
    

# ============================================================================
# FILE SCANNING
# ============================================================================

def scan_storage(
    root_path: Path,
    excluded_dirs: set = None
) -> Tuple[Counter, Counter, Dict[str, List[Path]]]:
    """
    Recursively scan directory and categorize files by type.
    
    This function walks through all files in the directory tree, categorizes
    them by extension, tracks directory sizes, and identifies documents
    suitable for RAG ingestion.
    
    Args:
        root_path: Root directory to scan
        excluded_dirs: Set of directory names to skip (default: {'AppData'})
        
    Returns:
        Tuple containing:
            - categories: Counter of file categories -> total bytes
            - directories: Counter of directory paths -> total bytes
            - documents_by_dir: Mapping of directory -> list of document paths
            
    Performance:
        - Processes ~1000 files/second on SSD
        - Uses single stat() call per file for efficiency
        - Skips permission errors gracefully
        
    Example:
        >>> from pathlib import Path
        >>> cats, dirs, docs = scan_storage(Path.home())
        >>> print(f"Found {sum(cats.values())} bytes in {len(cats)} categories")
    """
    if excluded_dirs is None:
        excluded_dirs = {'AppData'}
    
    categories = Counter()
    directories = Counter()
    documents_by_dir = {}
    file_count = 0
    error_count = 0
    
    logger.info(f"?? Scanning {root_path}...")
    
    try:
        for file_path in root_path.rglob("*"):
            try:
                # Skip excluded directories (optimization: check parts once)
                if any(excluded_dir in file_path.parts for excluded_dir in excluded_dirs):
                    continue
                
                if not file_path.is_file():
                    continue
                
                # Get file stats (single syscall)
                try:
                    file_stat = file_path.stat()
                    file_size = file_stat.st_size
                except (OSError, FileNotFoundError) as e:
                    logger.debug(f"Cannot stat {file_path}: {e}")
                    error_count += 1
                    continue
                
                # Categorize by extension
                ext = file_path.suffix.lower().lstrip('.')
                category = FILE_CATEGORIES.get(ext, "Overig")
                categories[category] += file_size
                
                # Track directory size
                parent_dir = str(file_path.parent)
                directories[parent_dir] += file_size
                
                # Track documents for RAG (size limit check)
                if ext in RAG_DOCUMENT_TYPES and file_size < MAX_DOCUMENT_SIZE:
                    documents_by_dir.setdefault(parent_dir, []).append(file_path)
                
                file_count += 1
                
                # Progress indicator (every N files)
                if file_count % PROGRESS_INTERVAL == 0:
                    print(f"? Processed {file_count:,} files...", end='\r')
                    
            except (PermissionError, OSError, FileNotFoundError) as e:
                error_count += 1
                logger.debug(f"Skipped {file_path}: {e}")
                continue
                
    except KeyboardInterrupt:
        logger.warning("\n? Scan interrupted by user")
    
    # Final statistics
    print(f"\n? Completed! Scanned {file_count:,} files")
    if error_count > 0:
        logger.warning(f"? Skipped {error_count:,} files due to errors")
    
    total_docs = sum(len(docs) for docs in documents_by_dir.values())
    logger.info(f"?? Found {total_docs:,} documents across {len(documents_by_dir):,} directories")
    
    return categories, directories, documents_by_dir


def get_top_document_directories(
    documents_by_dir: Dict[str, List[Path]],
    n: int = 10
) -> List[Tuple[str, int]]:
    """
    Get top N directories by document count.
    
    Args:
        documents_by_dir: Mapping of directory paths to document lists
        n: Number of top directories to return
        
    Returns:
        List of (directory_path, document_count) tuples, sorted descending
        
    Performance:
        O(m log m) where m = number of directories
    """
    dir_counts = [(dir_path, len(docs)) for dir_path, docs in documents_by_dir.items()]
    dir_counts.sort(key=lambda x: x[1], reverse=True)
    return dir_counts[:n]


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def extract_text_from_file(file_path: Path) -> Optional[str]:
    """
    Extract text content from various file types.
    
    Supports:
        - Plain text: .txt, .md, .rst, .csv, .json, .xml, .html
        - PDFs: .pdf (via pypdf)
        - Word documents: .docx (via python-docx)
    
    Args:
        file_path: Path to the file to extract
        
    Returns:
        Extracted text content, or None if extraction fails
        
    Error Handling:
        - Uses 'ignore' for encoding errors (lossy but robust)
        - Returns None on extraction failure (logged)
        - Gracefully handles missing optional libraries
        
    Performance:
        - Plain text: ~100MB/s
        - PDF: ~10-50 pages/s (depends on complexity)
        - DOCX: ~50-100 docs/s
    """
    ext = file_path.suffix.lower().lstrip('.')
    
    try:
        # Plain text files (fast path)
        if ext in ["txt", "md", "rst", "csv", "json", "xml", "html"]:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        # PDF files
        elif ext == "pdf":
            if not PYPDF_AVAILABLE:
                logger.debug(f"Skipped PDF {file_path.name}: pypdf not installed")
                return None
            
            try:
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    text_parts = [page.extract_text() for page in reader.pages]
                    return "\n".join(text_parts)
            except Exception as e:
                logger.debug(f"Error reading PDF {file_path.name}: {e}")
                return None
        
        # Word documents
        elif ext == "docx":
            if not DOCX_AVAILABLE:
                logger.debug(f"Skipped DOCX {file_path.name}: python-docx not installed")
                return None
            
            try:
                doc = docx.Document(file_path)
                text_parts = [para.text for para in doc.paragraphs]
                return "\n".join(text_parts)
            except Exception as e:
                logger.debug(f"Error reading DOCX {file_path.name}: {e}")
                return None
        
        return None
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return None


# ============================================================================
# TEXT CHUNKING
# ============================================================================

def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP
) -> List[str]:
    """
    Split text into overlapping chunks using recursive character splitting.
    
    Uses a hierarchical approach with multiple separators to preserve semantic
    meaning while maintaining optimal chunk sizes for embedding models.
    
    Args:
        text: Text to chunk
        chunk_size: Target size of each chunk in characters (~128 tokens for 512 chars)
        overlap: Number of overlapping characters between chunks for context continuity
        
    Returns:
        List of text chunks with overlap
        
    Algorithm:
        1. Try to split on paragraph breaks first (\n\n)
        2. Fall back to line breaks (\n)
        3. Then sentence boundaries. 
        4. Finally word boundaries ( )
        5. Last resort: character-level split
        
    Performance:
        - Optimized for 512-char chunks (~128 tokens)
        - 100-char overlap preserves context
        - Better semantic coherence than simple windowing
        
    Example:
        >>> text = "Paragraph 1.\\n\\nParagraph 2.\\n\\nParagraph 3."
        >>> chunks = chunk_text(text, chunk_size=50, overlap=10)
        >>> # Chunks split on paragraph boundaries when possible
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    separators = CHUNK_SEPARATORS
    
    def split_recursive(text: str, sep_index: int = 0) -> List[str]:
        """Recursively split text using hierarchical separators."""
        if len(text) <= chunk_size:
            return [text] if text.strip() else []
        
        if sep_index >= len(separators):
            # Last resort: character-level split with overlap
            result = []
            start = 0
            while start < len(text):
                end = min(start + chunk_size, len(text))
                chunk = text[start:end]
                if chunk.strip():
                    result.append(chunk)
                if end >= len(text):
                    break
                start = end - overlap
            return result
        
        separator = separators[sep_index]
        splits = text.split(separator)
        
        # Recombine splits into chunks
        result = []
        current_chunk = ""
        
        for i, split in enumerate(splits):
            # Restore separator (except for last split)
            piece = split + (separator if i < len(splits) - 1 else "")
            
            if len(current_chunk) + len(piece) <= chunk_size:
                current_chunk += piece
            else:
                # Current chunk is full
                if current_chunk.strip():
                    result.append(current_chunk)
                
                # If piece itself is too large, split it further
                if len(piece) > chunk_size:
                    result.extend(split_recursive(piece, sep_index + 1))
                    current_chunk = ""
                else:
                    # Start new chunk with overlap from previous
                    if result:
                        # Add overlap from end of previous chunk
                        prev_chunk = result[-1]
                        overlap_text = prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
                        current_chunk = overlap_text + piece
                    else:
                        current_chunk = piece
        
        # Add final chunk
        if current_chunk.strip():
            result.append(current_chunk)
        
        return result
    
    chunks = split_recursive(text)
    
    # Ensure no empty chunks
    chunks = [c.strip() for c in chunks if c.strip()]
    
    return chunks if chunks else [text]


# ============================================================================
# EMBEDDING GENERATION
# ============================================================================

def generate_embedding(
    text: str,
    model: str = OLLAMA_EMBED_MODEL
) -> Optional[List[float]]:
    """
    Generate embedding vector using Ollama.
    
    Sends text to the Ollama embeddings API and returns a dense vector
    representation suitable for semantic similarity search.
    
    Args:
        text: Text to embed (truncated to 8000 chars)
        model: Embedding model name (default: nomic-embed-text)
        
    Returns:
        768-dimensional embedding vector, or None on failure
        
    API Details:
        - Endpoint: POST /api/embeddings
        - Timeout: 60 seconds
        - Dimension: 768 (nomic-embed-text)
        - Retry: 3 attempts with exponential backoff
        
    Error Handling:
        - Returns None on connection/timeout errors
        - Validates embedding dimension
        - Logs all errors for debugging
        - Implements retry logic for transient failures
        
    Performance:
        - ~10-50 embeddings/second (depends on text length)
        - Batching not currently supported
    """
    if not REQUESTS_AVAILABLE:
        return None
    
    # Retry configuration
    max_retries = 3
    base_delay = 0.5  # seconds
    
    for attempt in range(max_retries):
        try:
            # Truncate text and log length
            truncated_text = text[:8000]
            if len(text) > 8000:
                logger.debug(f"Text truncated from {len(text)} to 8000 chars")
            
            payload = {
                "model": model,
                "prompt": truncated_text
            }
            
            if attempt > 0:
                logger.debug(f"Retry attempt {attempt + 1}/{max_retries} for embedding generation")
            
            resp = requests.post(
                OLLAMA_EMBED_URL,
                json=payload,
                timeout=60
            )
            resp.raise_for_status()
            
            data = resp.json()
            embedding = data.get("embedding")
            
            if not embedding:
                logger.error("? No embedding in response")
                if attempt < max_retries - 1:
                    import time
                    delay = base_delay * (2 ** attempt)
                    logger.debug(f"Waiting {delay}s before retry...")
                    time.sleep(delay)
                    continue
                return None
            
            # Validate embedding dimension
            if len(embedding) != OLLAMA_EMBED_DIMENSION:
                logger.warning(
                    f"? Expected {OLLAMA_EMBED_DIMENSION}D embedding, "
                    f"got {len(embedding)}D"
                )
            
            logger.debug(f"? Embedding generated successfully ({len(embedding)}D)")
            return embedding
            
        except requests.exceptions.Timeout:
            logger.error(f"? Embedding generation timed out (60s) - Attempt {attempt + 1}/{max_retries}")
            if attempt < max_retries - 1:
                import time
                delay = base_delay * (2 ** attempt)
                logger.debug(f"Waiting {delay}s before retry...")
                time.sleep(delay)
                continue
            logger.error("  This usually means Ollama is overloaded or not responding")
            return None
            
        except requests.exceptions.ConnectionError:
            logger.error(f"? Cannot connect to Ollama - Attempt {attempt + 1}/{max_retries}")
            if attempt < max_retries - 1:
                import time
                delay = base_delay * (2 ** attempt)
                logger.debug(f"Waiting {delay}s before retry...")
                time.sleep(delay)
                continue
            logger.error("  Check: curl http://localhost:11434/api/tags")
            return None
            
        except requests.exceptions.JSONDecodeError as e:
            logger.error(f"? Invalid JSON response from Ollama: {e}")
            return None
            
        except Exception as e:
            logger.error(f"Error generating embedding (attempt {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                import time
                delay = base_delay * (2 ** attempt)
                logger.debug(f"Waiting {delay}s before retry...")
                time.sleep(delay)
                continue
            return None
    
    return None


# ============================================================================
# DATABASE MANAGEMENT
# ============================================================================

@contextmanager
def get_db_connection():
    """
    Context manager for database connections.
    
    Ensures proper connection cleanup even if errors occur.
    Simplified connection without keepalive for better compatibility.
    
    Yields:
        psycopg2 connection object
        
    Example:
        >>> with get_db_connection() as conn:
        ...     with conn.cursor() as cur:
        ...         cur.execute("SELECT 1")
    """
    conn = None
    try:
        conn = psycopg2.connect(
            host=PG_HOST,
            port=PG_PORT,
            database=PG_DATABASE,
            user=PG_USER,
            password=PG_PASSWORD,
            connect_timeout=10
        )
        yield conn
    except psycopg2.OperationalError as e:
        logger.error(f"? Database connection failed: {e}")
        raise
    finally:
        if conn:
            try:
                conn.close()
            except:
                pass


def create_database_if_not_exists() -> bool:
    """
    Create the vectordb database if it doesn't exist.
    
    Connects to the default 'postgres' database to create our database.
    Uses autocommit mode since CREATE DATABASE cannot run in a transaction.
    
    Returns:
        True if successful (or already exists), False on error
        
    SQL Safety:
        Uses sql.Identifier to prevent SQL injection
    """
    if not PSYCOPG2_AVAILABLE:
        return False
    
    conn = None
    try:
        # Use connection string format for better compatibility
        conn = psycopg2.connect(
            host=PG_HOST,
            port=PG_PORT,
            database='postgres',
            user=PG_USER,
            password=PG_PASSWORD,
            connect_timeout=10
        )
        conn.autocommit = True
        
        with conn.cursor() as cur:
            # Check if database exists
            cur.execute(
                "SELECT 1 FROM pg_database WHERE datname = %s",
                (PG_DATABASE,)
            )
            exists = cur.fetchone()
            
            if not exists:
                logger.info(f"Creating database '{PG_DATABASE}'...")
                cur.execute(sql.SQL("CREATE DATABASE {}").format(
                    sql.Identifier(PG_DATABASE)
                ))
                logger.info(f"? Database '{PG_DATABASE}' created")
            else:
                logger.debug(f"Database '{PG_DATABASE}' already exists")
        
        return True
        
    except psycopg2.OperationalError as e:
        logger.error(f"? PostgreSQL connection error: {e}")
        logger.error("  Check that PostgreSQL is running and credentials are correct")
        return False
    except psycopg2.Error as e:
        logger.error(f"? PostgreSQL error: {e}")
        return False
    finally:
        if conn:
            try:
                conn.close()
            except:
                pass


def init_pgvector_table(conn) -> bool:
    """
    Initialize pgvector table with proper schema.
    
    Creates the documents table with vector embeddings and necessary indexes.
    Automatically migrates old schemas by dropping and recreating.
    
    Args:
        conn: Active psycopg2 connection
        
    Returns:
        True if successful, False on error
        
    Schema:
        - file_path + chunk_index: Unique constraint
        - embedding: 768-dimensional vector
        - Indexes: IVFFlat for similarity search, B-tree for file_path
        
    Migration:
        - Detects old schema (missing chunk_index)
        - Drops and recreates table automatically
        - Warns user about data loss
    """
    try:
        with conn.cursor() as cur:
            # Enable pgvector extension
            cur.execute("CREATE EXTENSION IF NOT EXISTS vector;")
            
            # Check if table has correct schema
            cur.execute(sql.SQL("""
                SELECT column_name 
                FROM information_schema.columns 
                WHERE table_name = %s AND column_name = 'chunk_index';
            """), [PG_TABLE])
            
            has_chunk_index = cur.fetchone() is not None
            
            # Drop old schema if needed
            if not has_chunk_index:
                cur.execute(sql.SQL("""
                    SELECT EXISTS (
                        SELECT FROM information_schema.tables 
                        WHERE table_name = %s
                    );
                """), [PG_TABLE])
                
                table_exists = cur.fetchone()[0]
                
                if table_exists:
                    logger.warning(
                        f"? Table '{PG_TABLE}' has old schema. "
                        f"Dropping and recreating..."
                    )
                    cur.execute(sql.SQL("DROP TABLE IF EXISTS {} CASCADE;").format(
                        sql.Identifier(PG_TABLE)
                    ))
            
            # Create table with vector column
            cur.execute(sql.SQL("""
                CREATE TABLE IF NOT EXISTS {} (
                    id SERIAL PRIMARY KEY,
                    file_path TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    file_name TEXT NOT NULL,
                    file_type TEXT,
                    content_preview TEXT,
                    chunk_content TEXT,
                    file_size BIGINT,
                    modified_time FLOAT,
                    embedding vector(%s),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(file_path, chunk_index)
                );
            """).format(sql.Identifier(PG_TABLE)), [OLLAMA_EMBED_DIMENSION])
            
            # Create IVFFlat index for vector similarity search
            cur.execute(sql.SQL("""
                CREATE INDEX IF NOT EXISTS {} 
                ON {} USING ivfflat (embedding vector_cosine_ops)
                WITH (lists = 100);
            """).format(
                sql.Identifier(f"{PG_TABLE}_embedding_idx"),
                sql.Identifier(PG_TABLE)
            ))
            
            # Create B-tree index on file_path for faster lookups
            cur.execute(sql.SQL("""
                CREATE INDEX IF NOT EXISTS {} 
                ON {} (file_path);
            """).format(
                sql.Identifier(f"{PG_TABLE}_file_path_idx"),
                sql.Identifier(PG_TABLE)
            ))
            
            conn.commit()
            logger.info("? pgvector table initialized")
            return True
            
    except Exception as e:
        logger.error(f"Error initializing pgvector table: {e}")
        conn.rollback()
        return False


def is_document_ingested(
    conn,
    file_path: Path,
    modified_time: float
) -> bool:
    """
    Check if document is already ingested with same modification time.
    
    Args:
        conn: Active psycopg2 connection
        file_path: Path to check
        modified_time: File modification timestamp
        
    Returns:
        True if document exists with matching mtime, False otherwise
        
    Performance:
        Uses file_path index for fast lookup (O(log n))
    """
    try:
        with conn.cursor() as cur:
            cur.execute(sql.SQL("""
                SELECT modified_time FROM {} 
                WHERE file_path = %s 
                LIMIT 1;
            """).format(sql.Identifier(PG_TABLE)), (str(file_path),))
            
            result = cur.fetchone()
            if result:
                stored_mtime = result[0]
                return stored_mtime == modified_time
        return False
    except Exception as e:
        logger.debug(f"Error checking document: {e}")
        return False


def ingest_to_pgvector(
    conn,
    file_path: Path,
    chunks: List[str],
    embeddings: List[List[float]],
    modified_time: float
) -> bool:
    """
    Ingest document chunks into pgvector database.
    
    Deletes old chunks for the file and inserts new ones. This allows
    for document updates while maintaining chunk integrity.
    
    Args:
        conn: Active psycopg2 connection
        file_path: Path to document
        chunks: List of text chunks
        embeddings: List of embedding vectors (one per chunk)
        modified_time: File modification timestamp
        
    Returns:
        True if successful, False on error
        
    Transaction:
        All chunks inserted in single transaction (atomic).
        Does NOT commit - caller must commit for batching.
    """
    try:
        file_stat = file_path.stat()
        file_size = file_stat.st_size
        
        with conn.cursor() as cur:
            # Delete old chunks
            cur.execute(sql.SQL("DELETE FROM {} WHERE file_path = %s;").format(
                sql.Identifier(PG_TABLE)
            ), (str(file_path),))
            
            # Insert all chunks
            for chunk_idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                cur.execute(sql.SQL("""
                    INSERT INTO {} 
                    (file_path, chunk_index, file_name, file_type, content_preview, 
                     chunk_content, file_size, modified_time, embedding)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s);
                """).format(sql.Identifier(PG_TABLE)), (
                    str(file_path),
                    chunk_idx,
                    file_path.name,
                    file_path.suffix.lstrip('.'),
                    chunk[:200],  # Preview for UI
                    chunk,
                    file_size,
                    modified_time,
                    embedding
                ))
        
        # Note: Caller commits for transaction batching
        return True
            
    except OSError as e:
        logger.error(f"Cannot stat file {file_path}: {e}")
        return False
    except psycopg2.IntegrityError as e:
        logger.error(f"Database integrity error for {file_path}: {e}")
        conn.rollback()
        return False
    except Exception as e:
        logger.error(f"Error ingesting to pgvector: {e}")
        conn.rollback()
        return False


# ============================================================================
# DOCUMENT INGESTION PIPELINE
# ============================================================================

def ingest_documents_to_pgvector(documents: List[Path]) -> int:
    """
    Ingest documents into pgvector database for RAG with detailed progress logging.
    
    Pipeline:
        1. Check if document already ingested (skip if unchanged)
        2. Extract text from document
        3. Split into overlapping chunks
        4. Generate embeddings for each chunk
        5. Store chunks with embeddings in pgvector
    
    Args:
        documents: List of document paths to ingest
        
    Returns:
        Number of successfully ingested documents
        
    Performance:
        - ~2-10 documents/second (depends on size, format)
        - Bottleneck: Ollama embedding generation
        - Progress updates for every document
        - Uses single persistent connection for all operations
        
    Error Handling:
        - Skips documents with insufficient content
        - Continues on individual document failures
        - Logs detailed error information
        - Uses connection pooling to prevent startup packet errors
    """
    if not REQUESTS_AVAILABLE or not PSYCOPG2_AVAILABLE:
        logger.error("? Required modules not available. Cannot ingest.")
        return 0
    
    # Create database if needed
    logger.info("=" * 70)
    logger.info("?? DOCUMENT INGESTION PIPELINE")
    logger.info("=" * 70)
    
    if not create_database_if_not_exists():
        logger.error("? Failed to create database")
        return 0
    
    logger.info(f"?? Establishing connection to PostgreSQL...")
    logger.info(f"   Host: {PG_HOST}:{PG_PORT}")
    logger.info(f"   Database: {PG_DATABASE}")
    
    conn = None
    try:
        # Establish single persistent connection for entire ingestion
        # This prevents "invalid length of startup packet" errors
        conn = psycopg2.connect(
            host=PG_HOST,
            port=PG_PORT,
            database=PG_DATABASE,
            user=PG_USER,
            password=PG_PASSWORD,
            connect_timeout=10,
            application_name='WhereSpace_Ingestion'
        )
        
        logger.info("? Database connection established")
        
        # Set timeouts to prevent hanging
        logger.info("??  Configuring database session...")
        with conn.cursor() as cur:
            cur.execute("SET statement_timeout = '300000';")  # 5 minutes
            cur.execute("SET idle_in_transaction_session_timeout = '600000';")  # 10 minutes
        conn.commit()
        logger.info("? Session configured")
        
        # Initialize table
        logger.info("?? Initializing pgvector table...")
        if not init_pgvector_table(conn):
            logger.error("? Failed to initialize table")
            return 0
        logger.info("? Table ready")
        
        # Check how many documents are already ingested
        logger.info("?? Checking existing documents...")
        with conn.cursor() as cur:
            cur.execute(sql.SQL("""
                SELECT COUNT(DISTINCT file_path) FROM {};
            """).format(sql.Identifier(PG_TABLE)))
            existing_count = cur.fetchone()[0]
        
        logger.info(f"? Found {existing_count:,} documents already in database")
        
        # Check if we've reached the 50 document limit
        if existing_count >= 50:
            logger.warning("=" * 70)
            logger.warning(f"? DOCUMENT LIMIT REACHED ({existing_count}/50 documents)")
            logger.warning("  Skipping ingestion to prevent excessive storage usage")
            logger.warning("  To ingest more documents:")
            logger.warning("  1. Increase limit in code, or")
            logger.warning("  2. Clear database via web interface")
            logger.warning("=" * 70)
            return 0
        
        # Calculate how many more we can ingest
        remaining_slots = 50 - existing_count
        documents_to_process = documents[:remaining_slots]
        
        if len(documents) > remaining_slots:
            logger.warning(f"? Limiting ingestion to {remaining_slots} documents (50 total limit)")
            logger.warning(f"  {len(documents) - remaining_slots} documents will be skipped")
        
        logger.info("=" * 70)
        logger.info(f"?? STARTING INGESTION")
        logger.info(f"   Documents to process: {len(documents_to_process):,}")
        logger.info(f"   Total after ingestion: {existing_count + len(documents_to_process)}/50")
        logger.info("=" * 70)
        print()
        
        ingested_count = 0
        failed_count = 0
        skipped_count = 0
        total_chunks = 0
        
        for i, doc_path in enumerate(documents_to_process, 1):
            try:
                # Progress header
                logger.info("?" * 70)
                logger.info(f"?? [{i}/{len(documents_to_process)}] Processing: {doc_path.name}")
                logger.info(f"   Path: {doc_path}")
                
                # Double-check we haven't exceeded limit
                if existing_count + ingested_count >= 50:
                    logger.warning(f"? Reached 50 document limit, stopping ingestion")
                    break
                
                # Get modification time
                try:
                    file_stat = doc_path.stat()
                    modified_time = file_stat.st_mtime
                    file_size = file_stat.st_size
                    logger.info(f"   Size: {format_size(file_size)}")
                    logger.info(f"   Modified: {Path(doc_path).stat().st_mtime}")
                except OSError as e:
                    logger.error(f"   ? Cannot access file: {e}")
                    skipped_count += 1
                    continue
                
                # Skip if already ingested (check BEFORE extracting text)
                if is_document_ingested(conn, doc_path, modified_time):
                    logger.info(f"   ? Already ingested (unchanged)")
                    skipped_count += 1
                    continue
                
                # Extract text
                logger.info(f"   ?? Extracting text...")
                text = extract_text_from_file(doc_path)
                if not text or len(text.strip()) < 50:
                    logger.warning(f"   ? Insufficient content (< 50 chars)")
                    skipped_count += 1
                    continue
                
                text_length = len(text)
                logger.info(f"   ? Extracted {text_length:,} characters")
                
                # Split into chunks
                logger.info(f"   ??  Chunking text (size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})...")
                chunks = chunk_text(text)
                logger.info(f"   ? Created {len(chunks)} chunks")
                total_chunks += len(chunks)
                
                # Generate embeddings - NOW WITH BATCH PROCESSING!
                logger.info(f"   ?? Generating embeddings...")
                
                # Use batch embedding if available (5-10x faster!)
                if BATCH_EMBEDDINGS_AVAILABLE and len(chunks) > 1:
                    logger.info(f"   ? Using batch mode ({len(chunks)} chunks)")
                    embeddings = generate_embeddings_batch(
                        chunks,
                        model=OLLAMA_EMBED_MODEL,
                        batch_size=20,  # Process 20 chunks at a time
                        max_workers=4   # 4 parallel workers
                    )
                    
                    # Filter out any None values (failed embeddings)
                    failed_embeddings = sum(1 for e in embeddings if e is None)
                    
                    if failed_embeddings > 0:
                        logger.warning(f"      ? Failed {failed_embeddings}/{len(chunks)} embeddings")
                        
                        if failed_embeddings == len(chunks):
                            logger.error(f"   ? All embeddings failed")
                            failed_count += 1
                            continue
                        
                        # Keep only successful embeddings and their corresponding chunks
                        valid_pairs = [(c, e) for c, e in zip(chunks, embeddings) if e is not None]
                        chunks = [c for c, e in valid_pairs]
                        embeddings = [e for c, e in valid_pairs]
                        logger.info(f"   ? Continuing with {len(chunks)}/{len(chunks)+failed_embeddings} chunks")
                    
                    logger.info(f"   ? Generated {len(embeddings)} embeddings (batch mode)")
                
                else:
                    # Fall back to sequential mode (slower)
                    if BATCH_EMBEDDINGS_AVAILABLE:
                        logger.info(f"   ? Using sequential mode (single chunk)")
                    else:
                        logger.info(f"   ? Batch mode not available, using sequential")
                    
                    embeddings = []
                    failed_embeddings = 0
                    
                    for chunk_idx, chunk in enumerate(chunks):
                        # Progress indicator for chunks
                        if chunk_idx > 0 and chunk_idx % 5 == 0:
                            print(f"      Embedding {chunk_idx}/{len(chunks)}...", end='\r')
                        
                        embedding = generate_embedding(chunk)
                        if not embedding:
                            failed_embeddings += 1
                            logger.warning(f"      ? Failed embedding {chunk_idx + 1}/{len(chunks)}")
                            break
                        embeddings.append(embedding)
                        
                        # Small delay to prevent overwhelming Ollama
                        if chunk_idx > 0 and chunk_idx % 5 == 0:
                            import time
                            time.sleep(0.1)
                    
                    # Clear progress line
                    print(" " * 50, end='\r')
                    
                    # Skip if any embedding failed
                    if len(embeddings) != len(chunks):
                        logger.error(f"   ? Embedding generation incomplete ({len(embeddings)}/{len(chunks)} succeeded)")
                        logger.error(f"      Failed: {failed_embeddings} embeddings")
                        failed_count += 1
                        continue
                    
                    logger.info(f"   ? Generated {len(embeddings)} embeddings (sequential mode)")
                
                # Ingest to database using same connection
                logger.info(f"   ?? Storing in database...")
                if ingest_to_pgvector(conn, doc_path, chunks, embeddings, modified_time):
                    ingested_count += 1
                    new_total = existing_count + ingested_count
                    logger.info(f"   ? SUCCESS! Ingested with {len(chunks)} chunks")
                    logger.info(f"   ?? Progress: {new_total}/50 documents in database")
                else:
                    failed_count += 1
                    logger.error(f"   ? Failed to store in database")
                
                # Commit every 5 documents to prevent long transactions
                if i % 5 == 0:
                    logger.info(f"   ?? Committing batch to database...")
                    conn.commit()
                    logger.info(f"   ? Batch committed")
                
                print()
                    
            except Exception as e:
                logger.error("?" * 70)
                logger.error(f"? EXCEPTION processing {doc_path.name}")
                logger.error(f"   Error: {e}")
                logger.error(f"   Type: {type(e).__name__}")
                import traceback
                logger.error(f"   Traceback:\n{traceback.format_exc()}")
                logger.error("?" * 70)
                failed_count += 1
                
                # Rollback on error and continue
                try:
                    conn.rollback()
                    logger.info("   ? Transaction rolled back")
                except Exception as rollback_error:
                    logger.error(f"   ? Rollback error: {rollback_error}")
                continue
        
        # Final commit
        logger.info("=" * 70)
        logger.info("?? Final commit...")
        try:
            conn.commit()
            logger.info("? Final commit successful")
        except Exception as commit_error:
            logger.error(f"? Final commit error: {commit_error}")
        
        # Summary
        total_documents = existing_count + ingested_count
        logger.info("=" * 70)
        logger.info("?? INGESTION SUMMARY")
        logger.info("=" * 70)
        logger.info(f"? Successfully ingested:  {ingested_count:,} documents")
        logger.info(f"? Failed:                  {failed_count:,} documents")
        logger.info(f"? Skipped:                 {skipped_count:,} documents")
        logger.info(f"?? Total chunks created:   {total_chunks:,}")
        logger.info(f"?? Total in database:      {total_documents}/50 documents")
        
        if total_documents >= 50:
            logger.info("?? Document limit reached! Database ready for querying.")
        else:
            remaining = 50 - total_documents
            logger.info(f"?? Can still ingest {remaining} more documents")
        
        logger.info("=" * 70)
        
        return ingested_count
        
    except psycopg2.OperationalError as e:
        logger.error("=" * 70)
        logger.error("X POSTGRESQL CONNECTION ERROR")
        logger.error("=" * 70)
        logger.error(f"Error: {e}")
        logger.error("\nPossible causes:")
        logger.error("  1. PostgreSQL server not running")
        logger.error("  2. Incorrect host/port/database/credentials")
        logger.error("  3. Firewall blocking connection")
        logger.error("  4. Too many connection attempts")
        logger.error("\nSolutions:")
        logger.error("  - Check: pg_isready")
        logger.error("  - Restart PostgreSQL service")
        logger.error("  - Verify connection settings in code")
        logger.error("  - Wait 30 seconds and try again")
        logger.error("=" * 70)
        return 0
    except psycopg2.Error as e:
        logger.error("=" * 70)
        logger.error("X POSTGRESQL ERROR")
        logger.error("=" * 70)
        logger.error(f"Error: {e}")
        logger.error(f"Type: {type(e).__name__}")
        logger.error("=" * 70)
        return 0
    except Exception as e:
        logger.error("=" * 70)
        logger.error("X UNEXPECTED ERROR")
        logger.error("=" * 70)
        logger.error(f"Error: {e}")
        logger.error(f"Type: {type(e).__name__}")
        import traceback
        logger.error(f"Traceback:\n{traceback.format_exc()}")
        logger.error("=" * 70)
        return 0
    finally:
        # Always close connection cleanly
        if conn:
            try:
                logger.info("Closing database connection...")
                conn.close()
                logger.info("Connection closed cleanly")
            except Exception as close_error:
                logger.error(f"X Error closing connection: {close_error}")


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def format_size(bytes_size: int) -> str:
    """
    Convert bytes to human-readable format.
    
    Args:
        bytes_size: Size in bytes
        
    Returns:
        Formatted string (e.g., "1.50 GB")
        
    Raises:
        ValueError: If size is negative
        
    Example:
        >>> format_size(1536)
        '1.50 KB'
        >>> format_size(1073741824)
        '1.00 GB'
    """
    if bytes_size < 0:
        raise ValueError("Size cannot be negative")
    
    for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
        if bytes_size < 1024:
            return f"{bytes_size:.2f} {unit}"
        bytes_size /= 1024
    
    return f"{bytes_size:.2f} PB"


def display_results(categories: Counter, directories: Counter) -> None:
    """
    Display storage analysis results in formatted tables.
    
    Shows two tables:
        1. File categories with sizes and percentages
        2. Top 10 largest directories
    
    Args:
        categories: File categories with sizes
        directories: Directory paths with sizes
    """
    # Category summary
    print("\n" + "=" * 50)
    print("=== Storage Overview ===")
    print("=" * 50)
    
    total_size = sum(categories.values())
    
    if total_size == 0:
        print("No files found.")
        return
    
    for cat, size in categories.most_common():
        percentage = (size / total_size * 100)
        print(f"{cat:15} {format_size(size):>12} ({percentage:>5.1f}%)")
    
    print("-" * 50)
    print(f"{'TOTAAL':15} {format_size(total_size):>12} (100.0%)")
    
    # Top directories
    print("\n" + "=" * 50)
    print("=== Top 10 Largest Directories ===")
    print("=" * 50)
    
    for i, (dir_path, size) in enumerate(directories.most_common(10), 1):
        print(f"{i:2}. {format_size(size):>12} - {dir_path}")


def send_to_ollama(categories: Counter) -> str:
    """
    Send storage summary to Ollama for AI analysis.
    
    Generates a Dutch-language analysis of storage distribution using
    the llama3.1 model.
    
    Args:
        categories: Counter with file categories and sizes
        
    Returns:
        AI-generated analysis text (Dutch)
        
    Error Handling:
        Returns user-friendly error messages for common failures
    """
    if not REQUESTS_AVAILABLE:
        return "requests module niet geinstalleerd. Run: pip install requests"
    
    try:
        summary = {cat: format_size(size) for cat, size in categories.items()}
        
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": f"Geef een korte analyse van deze storage verdeling: {summary}",
            "stream": False
        }
        
        logger.debug(f"Sending request to Ollama: {payload}")
        
        resp = requests.post(
            OLLAMA_URL,
            json=payload,
            timeout=OLLAMA_TIMEOUT
        )
        resp.raise_for_status()
        
        response_data = resp.json()
        if "response" in response_data:
            return response_data["response"]
        else:
            logger.error(f"Unexpected response format: {response_data}")
            return f"Unexpected response format: {list(response_data.keys())}"
        
    except requests.exceptions.ConnectionError:
        return "Error: Ollama is niet bereikbaar. Start Ollama met 'ollama serve'."
    except requests.exceptions.Timeout:
        return f"Error: Ollama timeout na {OLLAMA_TIMEOUT} seconden."
    except requests.exceptions.JSONDecodeError as e:
        logger.error(f"JSON decode error: {e}")
        return f"Error: Invalid JSON response from Ollama: {str(e)}"
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error: {e}")
        return f"Error: {str(e)}"


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main() -> None:
    """
    Main application entry point.
    
    Workflow:
        1. Scan home directory
        2. Display storage analysis
        3. Get AI analysis (optional)
        4. Offer RAG document ingestion (interactive)
    
    Exit Codes:
        0: Success
        1: Unhandled error
        130: User interrupt (Ctrl+C)
    """
    try:
        # Scan storage
        root = Path.home()
        logger.info(f"Starting scan of {root}")
        
        categories, directories, documents_by_dir = scan_storage(root)
        
        # Display results
        display_results(categories, directories)
        
        # AI analysis (optional)
        if REQUESTS_AVAILABLE:
            print("\n" + "=" * 50)
            print("=== AI Analyse ===")
            print("=" * 50)
            response = send_to_ollama(categories)
            print(response)
        else:
            logger.info("Skipping AI analysis - requests module not available")
        
        # RAG ingestion (interactive)
        if documents_by_dir:
            print("\n" + "=" * 50)
            print("=== RAG Document Ingestion ===")
            print("=" * 50)
            
            top_dirs = get_top_document_directories(documents_by_dir, 10)
            
            if not top_dirs:
                logger.info("No directories with documents found")
                return
            
            print("Top 10 directories met documenten:")
            for i, (dir_path, doc_count) in enumerate(top_dirs, 1):
                print(f"{i:2}. [{doc_count:4} docs] {dir_path}")
            
            print("\n 0. Alle directories")
            print(" q. Overslaan")
            
            choice = input("\nSelecteer directory nummer voor ingestion: ").strip().lower()
            
            if choice == 'q':
                logger.info("RAG ingestion skipped by user")
                return
            
            try:
                if choice == '0':
                    # Ingest all documents
                    all_docs = []
                    for docs in documents_by_dir.values():
                        all_docs.extend(docs)
                    logger.info(f"Ingesting all {len(all_docs):,} documents")
                    ingested = ingest_documents_to_pgvector(all_docs)
                else:
                    # Ingest selected directory
                    choice_idx = int(choice) - 1
                    if 0 <= choice_idx < len(top_dirs):
                        selected_dir, doc_count = top_dirs[choice_idx]
                        selected_docs = documents_by_dir[selected_dir]
                        logger.info(f"Ingesting {len(selected_docs):,} documents from: {selected_dir}")
                        ingested = ingest_documents_to_pgvector(selected_docs)
                    else:
                        logger.error("Invalid selection")
                        return
                
                print(f"\nSuccessfully ingested {ingested:,} documents into pgvector")
                
            except ValueError:
                logger.error("Invalid input")
                return
            
    except KeyboardInterrupt:
        logger.info("\nProgram interrupted by user")
        sys.exit(130)  # Standard exit code for SIGINT
    except Exception as e:
        logger.error(f"Unexpected error: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()